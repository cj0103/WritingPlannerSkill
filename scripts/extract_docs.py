#!/usr/bin/env python3
"""批量提取 PDF/DOCX/DOC 文件文本内容，输出到指定目录"""

import sys, os, json, hashlib

SRC_DIR = r"D:\工作文档\09_金融科技专题"
OUT_DIR = r"D:\Ollama\workbuddy\WritingPlannerSkill\knowledge_base\_raw_texts"

def extract_pdf_pypdf2(path):
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    return "\n".join(p.extract_text() or "" for p in reader.pages)

def extract_pdf_pdfplumber(path):
    import pdfplumber
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n".join(texts)

def extract_docx(path):
    from docx import Document
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        if p.text.strip():
            paras.append(p.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            paras.append(" | ".join(cells))
    return "\n".join(paras)

def extract_doc(path):
    """Extract .doc via olefile + antiword or fallback"""
    import olefile
    try:
        ole = olefile.OleFileIO(path)
        # Try to extract WordDocument stream
        if ole.exists('WordDocument'):
            # Try reading the text via a simple approach
            try:
                stream = ole.openstream('1Table')
                data = stream.read()
                # Try extracting readable text fragments
                text = data.decode('utf-16-le', errors='ignore')
                # Filter non-printable chars
                import re
                text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
                # Heuristic: keep lines with mostly Chinese chars
                lines = [l.strip() for l in text.split('\n') if l.strip()]
                # Keep lines that have a reasonable mix of Chinese and length
                result = []
                for l in lines:
                    if len(l) >= 4 and any('\u4e00' <= c <= '\u9fff' for c in l):
                        result.append(l)
                if result:
                    return "\n".join(result)
            except:
                pass
        ole.close()
    except:
        pass
    return "[未能提取 .doc 文件文本内容，请手动转换]"

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    
    files = []
    for root, dirs, fnames in os.walk(SRC_DIR):
        for fn in fnames:
            ext = os.path.splitext(fn)[1].lower()
            if ext in ('.pdf', '.docx', '.doc', '.md', '.txt'):
                files.append(os.path.join(root, fn))
    
    results = {}
    for fpath in files:
        fn = os.path.basename(fpath)
        ext = os.path.splitext(fn)[1].lower()
        print(f"Processing: {fn} ...", end=" ", flush=True)
        
        text = ""
        try:
            if ext == '.pdf':
                try:
                    text = extract_pdf_pdfplumber(fpath)
                except:
                    text = extract_pdf_pypdf2(fpath)
            elif ext == '.docx':
                text = extract_docx(fpath)
            elif ext == '.doc':
                text = extract_doc(fpath)
            elif ext in ('.md', '.txt'):
                with open(fpath, 'r', encoding='utf-8') as f:
                    text = f.read()
        except Exception as e:
            text = f"[提取失败: {e}]"
        
        # Calculate hash
        with open(fpath, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()
        
        # Save raw text
        safe_name = os.path.splitext(fn)[0].replace(' ', '_') + '.txt'
        out_path = os.path.join(OUT_DIR, safe_name)
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        char_count = len(text)
        print(f"→ {char_count} chars, hash={file_hash[:8]}...")
        
        results[fn] = {
            "source_path": fpath,
            "raw_text_path": out_path,
            "file_hash": file_hash,
            "char_count": char_count
        }
    
    # Save manifest
    manifest_path = os.path.join(OUT_DIR, "_manifest.json")
    with open(manifest_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"\nDone! {len(results)} files extracted to {OUT_DIR}")
    print(f"Manifest saved to {manifest_path}")

if __name__ == '__main__':
    main()
